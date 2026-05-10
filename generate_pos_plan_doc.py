import os
from docx import Document
from docx.shared import Pt

def generate_plan():
    doc = Document()

    # Title
    title = doc.add_heading('POS System Integration Plan', 0)
    title.alignment = 1

    doc.add_paragraph('Business: Lowell Nails & Spa')
    doc.add_paragraph('Orchestration Layer: Michael (AI Receptionist)')

    # Section 1: Objectives
    doc.add_heading('1. Vision & Core Objectives', level=1)
    doc.add_paragraph(
        "The objective is to evolve the current scheduling and CRM engine into a full-scale Point of Sale (POS) ecosystem. "
        "This bridges the gap between 'Scheduled Intent' and 'Financial Completion'."
    )

    # Section 2: Phases
    doc.add_heading('2. Implementation Phases', level=1)
    
    doc.add_heading('Phase 1: Core Transactions (The Wallet)', level=2)
    bullets = [
        "Payment Gateway: Integration with Stripe Terminal for in-person card processing.",
        "Deterministic Ledger: Creation of 'transactions' and 'line_items' SQL tables.",
        "Receipt Engine: Automated digital receipts via Twilio SMS and email."
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_heading('Phase 2: Operations (The Back-Office)', level=2)
    bullets = [
        "Technician Payroll: Automated commission calculations (e.g., 60/40 splits) based on staff config.",
        "Gift Card Logic: Full lifecycle management (Sale, Redemption, Balance Checking).",
        "Inventory Management: Tracking retail products and professional backbar supplies."
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Section 3: Technical Design
    doc.add_heading('3. Technical Architecture', level=1)
    doc.add_paragraph("Database: PostgreSQL schema expansion to support relational financial data.")
    doc.add_paragraph("API: Node.js Express routes for /api/checkout and /api/payroll.")
    doc.add_paragraph("AI Tools: New WAT tools for Michael to handle pricing and gift card balance queries.")

    # Section 4: Rollout
    doc.add_heading('4. Rollout Strategy', level=1)
    doc.add_paragraph("1. Schema Migration & Ledger Implementation")
    doc.add_paragraph("2. Payment Gateway Sandbox Testing")
    doc.add_paragraph("3. Frontend Checkout UI Integration")
    doc.add_paragraph("4. Automated Reporting & Payroll Dashboard")

    # Save
    output_path = "POS_Integration_Plan.docx"
    doc.save(output_path)
    print(f"✅ Success: Document generated at {os.path.abspath(output_path)}")

if __name__ == "__main__":
    try:
        generate_plan()
    except ImportError:
        print("❌ Error: python-docx not found. Run 'pip install python-docx'")